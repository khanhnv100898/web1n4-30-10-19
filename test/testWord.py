from docx import Document
from docx.shared import Inches
from pymongo import MongoClient
from bson import ObjectId


uri = "mongodb+srv://admin:admin@ttw-xlquo.mongodb.net/admin?retryWrites=true&w=majority"

client = MongoClient(uri)

# creat data
ttw_db = client.ttw_app

# creat collection
Order = ttw_db["orders"]
Login = ttw_db["logins"]


order = Order.find_one({"_id":ObjectId("5db931f9959ae0ff39036ab9")})
all_product =order['product_id']
user = Login.find_one({"_id":ObjectId("5da90b4fea26fac9ecd9a3d3")})


# 
document = Document()

document.add_heading('Hóa đơn bán hàng', 0)

document.add_heading("Tên khách hàng:  " + user['fullname'], level=1)

index = 1
for i in all_product:
    name = i['name']
    price = str(i['price'])

    s = "Sản phẩm " + str(index) +" :  " + name + "             " + price+"$"
    document.add_paragraph(s, style='Intense Quote')
    index += 1


document.add_paragraph("Địa chỉ:  " + order['address'] , style='Intense Quote')
document.add_paragraph("Phí ship:  " + str(order['ship_fee']) , style='Intense Quote')
document.add_paragraph("Tổng tiền hàng : " + str(order['order_fee']) , style='Intense Quote')
totalPrice = int(order['ship_fee']) + int(order['order_fee'])
document.add_paragraph("Tổng thanh toán:  " + str(totalPrice) , style='Intense Quote')


document.add_page_break()

document.save('demoHoaDonBanHang.docx')