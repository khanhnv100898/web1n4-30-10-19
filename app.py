from bson import ObjectId
from flask import Flask, render_template, request, redirect, session
from gmail import *
from datetime import datetime
import pyexcel
from pymongo import MongoClient
from bson import ObjectId
from docx import Document
from docx.shared import Inches

# uri để kết nối đến database
uri = "mongodb+srv://admin:admin@ttw-xlquo.mongodb.net/admin?retryWrites=true&w=majority"

# thực hiện kết nối
client = MongoClient(uri)
# get data
ttw_db = client.ttw_app

# get collection
Login = ttw_db["logins"]
Order = ttw_db["orders"]
Product = ttw_db["products"]


app = Flask(__name__)
app.config['SECRET_KEY'] = 'sdasdđâsdasasd23423@#4'

# Trang chủ
@app.route('/')
def index():
    return render_template('index.html') # trả về trang index.html

# Search
# Title search
@app.route('/searcht/searchTitle', methods=['GET', 'POST'])
def searchT():
    if request.method == "GET":
        return render_template('search/searchName.html')
    elif request.method == "POST":
        form = request.form
        productName = form['productName'].upper()
        s = "^"+productName
        all_product = Product.find({"name": {"$regex": s}})
        c = int(all_product.count())
        if c > 0:
            return render_template('search/search.html', all_product=all_product)
        else:
            return render_template('search/searchNone.html', template=0)
            
# Hiển thị danh sách sản phẩm
@app.route('/search/<search>')
def search(search):
    session['search'] = search
    all_product_gender = Product.find({'product_gender': search})
    if all_product_gender is not None:
        return render_template('search/search.html', all_product=all_product_gender)
    # elif session['search'] == "Kids":
    elif search == "Kids":
        all_product_kids = Product.find({'product_kids': True})
        if all_product_kids is not None:
            return render_template('search/search.html', all_product=all_product_kids)

# Search Info Kids
@app.route('/search_kids_info')
def search_kids_info():
    product = Product.find({"product_kids": True})
    return render_template('search/search.html', all_product=product)

# Search bestselling
@app.route('/searchbestselling')
def search_best_selling():
    all_product = Product.find({"sold_count": {"$gt": 150}})
    if all_product is not None:
        return render_template('search/bestselling.html', all_product=all_product)

# Searcg theo mẫu
@app.route('/search_type/<search_type>')
def search_type(search_type):
    session['search_type'] = search_type
    all_product_type = Product.find({"product_type": search_type})
    if all_product_type is not None:
        return render_template('search/search.html', all_product=all_product_type)

# Chi tiết sản phẩm
@app.route('/product/detail/<id>')
def productDetail(id):
    product = Product.find_one({"_id": ObjectId(id)})
    view_ex = int(product['view'])
    view_new = view_ex + 1
    update = {"$set": {"view": view_new}}
    Product.update_one({"_id": ObjectId(id)}, update)
    return render_template('search/detail.html', detail_product=product)

# Đăng ký
@app.route('/sign_up', methods=['GET', 'POST'])
def signUp():
    if request.method == 'GET':
        return render_template('signup/signup-user.html')
    elif request.method == 'POST':
        form = request.form
        fullname = form['fullname']
        username = form['username']
        password = form['password']
        # level = form['level']
        email = form['email']
        phone = form['phone']
        address = form['address']
        id_person = form['id_person']
        balance = form['balance']

        found_login = Login.find({"username": username})
        c = int(found_login.count())
        if c > 0:
            return render_template('signup/signup-fail.html', template=1)
        else:
            new_login = {
                "fullname": fullname,
                "username": username,
                "password": password,
                "level": 2,
                "email": email,
                "phone": phone,
                "id_person": id_person,
                "address": address,
                "balance": balance,
            }
            Login.insert_one(new_login)

            found_login = Login.find_one({'username': username})
            gmail = GMail('ngovankhanh108@gmail.com', 'KhanhNgo108')
            html_content = '''
        Chúc mừng bạn đã tạo tài khoản thành công trên hệ thống.
        Cảm ơn quý khách đã sử dụng dịch vụ của chúng tôi!
        '''
            msg = Message('Tạo tài khoản thành công',
                          to=found_login['email'], html=html_content)
            gmail.send(msg)
            return redirect('/login')

# Đăng nhập phân quyền
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'GET':
        return render_template('logins/login-user.html')  # trả về html khi người dùng yêu cầu
    elif request.method == 'POST':
        form = request.form 
        login_username = form['username']
        login_password = form['password']
        login = Login.find_one({'username': login_username}) # tìm kiếm username có trong database
        if login is None:  # nếu không tìm thấy username
            return render_template('logins/login-fail.html') # thông báo đăng nhập không thành công
        else: 
            if login_password == login['password']: # so sánh password người dùng nhập với password trong database nếu trùng thì đăng nhập
                session['loggedin'] = True # lưu phiên đăng nhập
                session['level_loggedin'] = str((login['level'])) # lưu phiên thông tin level của người đăng nhập
                # phân quyền đăng nhập
                if session['level_loggedin'] == "2": # user level = 2, lưu vào phiên đăng nhập
                    session["user_id"] = str(login["_id"]) # lưu session id  = id của user đăng nhập
                    return render_template('user.html', login=login) # trả về html và truyền login sang html
                elif session['level_loggedin'] == "1":# level = 1 sẽ là admin
                    session["admin_id"] = str(login["_id"])
                    return render_template('admin.html', login=login)
                elif session['level_loggedin'] == "3": #level sẽ là shipper
                    session["shipper_id"] = str(login["_id"])
                    return render_template('shipper.html', login=login)
            else:
                return render_template('logins/login-fail.html') # nếu password sai, trả về Login Fail

# Đăng xuất
@app.route('/logout')
def logout():
    if "loggedin" in session: # nếu tồn tại phiên trong session
        if session['loggedin'] == False:  # đã đăng xuất thì trả về trang chủ
            return redirect('/')
        else:
            session['loggedin'] = False # phiên bằng false và kiểm tra level đang đăng nhập và thực hiện xóa phiên
            if session['level_loggedin'] == "1": # thực hiện kiểm tra phiên đăng nhập 
                del session['admin_id']  # thực hiện xóa session để dăng xuất tài khoản
                del session['level_loggedin']
            elif session['level_loggedin'] == "2":
                del session['user_id']
                del session['level_loggedin']
            elif session['level_loggedin'] == "3":
                del session['shipper_id']
                del session['level_loggedin']
            return redirect('/')
    else:
        return redirect('/') 

# Update thông tin đăng nhập
@app.route('/all_login/update/<id>', methods=['GET', 'POST']) # <id> sẽ được truyền vào bằng id của document(bản ghi)
def updateLogin(id): # thực hiện hàm update với 1 tham số id được truyền vào
    login = Login.find_one({"_id": ObjectId(id)}) # thực hiện tìm kiếm trường "_id" trong collection với tham số id được truyền vào 
    if request.method == 'GET': 
        return render_template('update_login.html', login=login) # trả về trang update truyền document vừa tìm được sang html
    elif request.method == 'POST': # nếu người dùng thực hiện gửi dữ liệu, thực hiện lấy dữ liệu
        form = request.form     # lấy dữ liệu từ form
        fullname = form['fullname'] # gán dữ liệu ở form vào biến
        username = form['username']
        password = form['password']
        level = form['level']

        # update database
        update_login = {"$set": {
            "fullname": fullname,
            "username": username,
            "password": password,
            "level": level,
        }}
        Login.update_one(login, update_login)
        return redirect('/all_login')

# Update thông tin đăng nhập User
@app.route('/update/<id>', methods=['GET', 'POST'])
def updateLoginUser(id): # hàm update thông tin của user
    login = Login.find_one({"_id": ObjectId(id)}) # tìm user đang đăng nhập
    if request.method == 'GET':
        return render_template('user/user-update.html', update_user=login)
    elif request.method == 'POST':
        # lấy dữ liệu từ form 
        form = request.form
        fullname = form['fullname']
        username = form['username']
        password = form['password']
        phone = form['phone']
        email = form['email']
        address = form['address']
        # level = form['level']

        # thực hiện update bằng pymongo
        update_login = {"$set": {
            "fullname": fullname,
            "username": username,
            "password": password,
            "level": 2,
            "phone": phone,
            "email": email,
            "address": address,
        }}
        Login.update_one(login, update_login)
        return redirect('/homeLogin')

# Giao diện Information
@app.route('/homeLogin') 
def homeLogin():        # hàm giao diện của tài khoản khi đăng nhập
    if "loggedin" in session:
        if session['loggedin'] == True: # khi có phiên đăng nhập tồn tại, thực hiện kiểm tra 
            if session['level_loggedin'] == "2": 
                return render_template('user.html') # trả về trang user nếu session == 2
            elif session['level_loggedin'] == "1":
                return render_template('admin.html')# trả về trang admin nếu session == 1
            elif session['level_loggedin'] == "3":
                return render_template('shipper.html')# trả về trang shipper nếu session == 3
        else:
            return redirect('/login')
    return redirect('/login')

# Giao diện chính cho User
@app.route('/user')
def user():
    if "user-loggedin" in session:
        if session['user-loggedin'] == True:
            return render_template('user.html')
        else:
            return redirect('/login')
    else:
        return redirect('/login')

# Thông tin đăng nhập
# User
@app.route('/user_information')
def userInformation():
    id = session["user_id"]
    login = Login.find_one({"_id": ObjectId(id)})
    if login is not None:
        return render_template('user/information.html', found_user=login)
    else:
        return 'Login is not found'

# Nạp tiền cho tài khoản User
@app.route('/add_money_user', methods=['GET', 'POST'])
def addMoneyUser(): # hàm nạp tiền cho tài khoản user
    user_login = Login.find_one({"_id": ObjectId(session["user_id"])}) # thực hiện tìm kiếm tài khoản đang đăng nhập với id lưu session
    if request.method == 'GET':
        return render_template('user/add-money.html', found_user=user_login)
    elif request.method == 'POST':
        form = request.form
        money = int(form['money']) # lấy số tiền user nhập muốn nạp
        ex_money = int(user_login['balance']) # lấy số tiền hiện có trong tài khoản của user
        balance = money + ex_money # tính total bằng số tiền cũ + số tiền user nhập

        # thực hiện update số tiền trong tài khoản
        add_money = {"$set": {
            "balance": balance,
        }}

        Login.update_one(user_login, add_money)
        return redirect('/add_money_user')

# Thêm sản phẩm vào giỏ hàng
@app.route('/add_order/<product_id>', methods=['GET', 'POST'])
def addOrder(product_id): # thực hiện hàm tạo đơn order, thêm sản phẩm vào giỏ hàng
    if "loggedin" in session: # nếu tồn tại phiên đăng nhập thì thực hiện
        if session['loggedin'] == True:
            if session['level_loggedin'] == "2": # nếu là người dùng vơi vai trò user mới thực hiện
                # tìm kiếm thông tin user
                found_user = Login.find_one({"_id": ObjectId(session["user_id"])})
                # tìm kiếm thông tin sản phẩm
                found_product = Product.find_one({"_id": ObjectId(product_id)})
                # tìm kiếm thông tin đơn order
                found_order = Order.find_one({"user_id": session['user_id'], "is_ordered": False, "status": "Shipper chưa nhận đơn"})
                # nếu tìm thấy sản phẩm thì tiếp tục thực hiện
                if found_product is not None:
                    # nếu đơn order đã tồn tại thực hiện thêm sản phẩm
                    if found_order is not None:
                        # lấy tổng tiền hiện tại của đơn order
                        ex_fee = found_order['order_fee']
                        # cập nhật tổng tiền = giá bán sản phẩm + tổng tiền hiện tại
                        order_fee = int(found_product['price'])+int(ex_fee)
                        # phí ship sẽ bằng 0,1 x tổng tiền
                        ship_fee = 0.1 * \
                            int(int(found_product['price'])+int(ex_fee))
                        # Tăng số lượng đã bán của sản phẩm
                        sold_count_ex = int(found_product['sold_count'])
                        sold_count_new = sold_count_ex + 1
                        update_product2 = {
                            "$set": {"sold_count": sold_count_new}}
                        Product.update_one(found_product, update_product2)
                        #  Cập nhật giá tiền mới
                        found_order_update_2 = {"$set": {
                            "order_time": datetime.now(),
                            "order_fee": order_fee,
                            "ship_fee": ship_fee,
                            # Thêm 1 sản phẩm vào giỏ hàng
                        }, "$addToSet": {"product_id": { # thông tin sản phẩm
                            "_id": found_product['_id'],
                            "image": found_product['image'],
                            "name": found_product['name'],
                            "description": found_product['description'],
                            "price": int(found_product['price']),
                            "status": found_product['status'],
                            "brand": found_product['brand'],
                            "product_type": found_product['product_type'],
                            "product_gender": found_product['product_gender'],
                            "product_kids": found_product['product_kids'],
                        }}}
                        Order.update_one(found_order, found_order_update_2)
                        b = session['search']
                        a = "/search/" + b
                        return redirect(a)
                    else: # trường hợp còn lại là chưa tồn tại đơn order, và sẽ tạo mới
                        # Tăng số lượng đã bán của sản phẩm
                        sold_count_ex = int(found_product['sold_count'])
                        sold_count_new = sold_count_ex + 1
                        update_product2 = {
                            "$set": {"sold_count": sold_count_new}}
                        Product.update_one(found_product, update_product2)
                        # Tạo 1 đơn mới
                        add_order = {
                            "user_id": session['user_id'], # lưu id người mua
                            "product_id": [{                # lưu thông tin sản phẩm
                                "_id": found_product['_id'],
                                "image": found_product['image'],
                                "name": found_product['name'],
                                "description": found_product['description'],
                                "price": int(found_product['price']),
                                "status": found_product['status'],
                                "brand": found_product['brand'],
                                "product_type": found_product['product_type'],
                                "product_gender": found_product['product_gender'],
                                "product_kids": found_product['product_kids'],
                            }],
                            "address": found_user['address'], # địa chỉ người mua
                            "order_time": datetime.now(), # thời gian gửi yêu cầu
                            "order_fee": int(found_product['price']), # tiền sản phẩm
                            "ship_fee": 0.1*int(found_product['price']), #phí ship
                            "is_ordered": False, # trạng thái order
                            "status": "Shipper chưa nhận đơn", # trạng thái của đơn hàng khi vận chuyển
                            "shipper_id": "null", # id của shipper nhận đơn
                        }
                        Order.insert_one(add_order)
                        
                        b = session['search']
                        a = "/search/" + b
                        return redirect(a)
                else:
                    return 'Not found'
            else:
                return redirect('/login')
        else:
            return redirect('/login')
    else:
        return redirect('/login')

# Hiển thị giỏ hàng
@app.route('/cart', methods=['GET', 'POST']) 
def cart(): #  hàm hiển thị giỏ hàng
    if "loggedin" in session: # nếu tồn tại phiên đăng nhập
        if session['loggedin'] == True:
            if session['level_loggedin'] == "2": # kiểm tra nếu level phiên đăng nhập == 2 là user thì thực hiện
                # tìm đơn order hiện tại của user
                found_order = Order.find_one({"user_id": session['user_id'], 'is_ordered': False, 'status': "Shipper chưa nhận đơn"})
                # nếu người dùng yêu cầu trang giỏ hàng thì sẽ trả về trang "cart.html"
                if request.method == 'GET':
                    if found_order is not None: # nếu order tồn tại thì thực hiện hiển thị
                        cart = found_order
                        order_product = cart['product_id']
                        return render_template('user/cart.html', cart=cart, order_product=order_product, template=1)
                    else: # nếu không tìm thấy đơn order nào sẽ trả về đơn hàng trống
                        return render_template('user/cart.html', template=0)
                # chức năng cập nhật thời gian yêu cầu nhận hàng, mặc định sẽ là giao hàng bất cứ lúc nào(chỉnh sửa thời gian muốn nhận hàng theo ý muốn)
                elif request.method == 'POST':
                    form = request.form
                    request_time = form['request_time']
                    # thực hiện lưu mới vào database
                    update = {"$set": {
                        'request_time': request_time,
                    }}
                    Order.update_one(found_order, update)
                return redirect('/cart')
            # nếu level loggedin == 3 là shipper thì không có quyền được vào trang giỏ hàng
            elif session['level_loggedin'] == "3":
                return redirect('/')
            # nếu level loggedin == 1 là 1 thì không có quyền được vào trang giỏ hàng
            elif session['level_loggedin'] == "1":
                return redirect('/')
        else:
            return redirect('/login')
    else:
        return redirect('/login')

# Xóa một sản phẩm khỏi giỏ hàng
@app.route('/delete_product/<product_id>')
def delete_product(product_id): # hàm xóa 1 sản phẩm có trong giỏ hàng
    # tìm đơn order hiện có của user đang đăng nhập
    order = Order.find_one({'user_id': session['user_id'], 'is_ordered': False, 'status': "Shipper chưa nhận đơn"})

    c = order['product_id']
    # nếu tồn tại đơn order
    if len(c) > 1:
        # thực hiện tìm thông tin sản phẩm muốn xóa = id sản phẩm
        product = Product.find_one({"_id": ObjectId(product_id)})
        # Giảm số lượng đã bán của sản phẩm
        sold_count_ex = int(product['sold_count'])
        sold_count_new = sold_count_ex - 1
        update_product2 = {"$set": {"sold_count": sold_count_new}}
        Product.update_one(product, update_product2)
        # Xóa 1 sản phẩm
        Order.update_one(
            # update đơn order hiện tại user muốn xóa sản phẩm
            {'_id': ObjectId(order['_id'])},
            # xóa 1 object khỏi trưởng product_id thuộc kiểu array , $pull là lấy ra
            {'$pull': {"product_id": {'_id': ObjectId(product_id)}}},
        )
        # cập nhật lại giá tiền khi xóa 1 sản phẩm
        order_fee_ex = int(order['order_fee'])
        product_price = int(product['price'])
        new_order_fee = int(order_fee_ex) - int(product_price)
        new_ship_fee = (0.1 * new_order_fee)
        # thực hiện cập nhật trên database
        Order.update_one(
            {'_id': ObjectId(order['_id'])},
            {'$set': {"order_fee": new_order_fee, "ship_fee": new_ship_fee}},
        )
        return redirect('/cart')
    # trường hợp đơn order tồn tại duy nhất 1 sản phẩm
    elif len(c) == 1:
        #thực hiện tìm kiếm thông tin sản phẩm bằng id
        product = Product.find_one({"_id": ObjectId(product_id)})
        # Giảm số lượng đã bán của sản phẩm
        sold_count_ex = int(product['sold_count'])
        sold_count_new = sold_count_ex - 1
        update_product2 = {"$set": {"sold_count": sold_count_new}}
        Product.update_one(product, update_product2)
        # thực hiện xóa document hiện tại khi chỉ còn 1 object(sản phẩm) trong trường "product_id"
        Order.delete_one(order)
        return redirect('/cart')

# Gửi yêu cầu mua đơn hàng
@app.route('/ordered/<order_id>')
def ordered(order_id): # hàm gửi yêu cầu nhận id của đơn order
    # tìm thông tin đơn order
    found_order = Order.find_one({"_id": ObjectId(order_id)})
    # tìm thông tin người dùng gửi yêu cầu
    found_user = Login.find_one({"_id": ObjectId(found_order['user_id'])})
    # nếu tồn tại đơn order
    if found_order is not None:
        # tổng tiền sẽ bằng phí sản phẩm + phí ship
        total = int(found_order['order_fee']) + int(found_order['ship_fee'])
        # lấy số tiền hiện có của user
        ex_balance = int(found_user['balance'])
        # nếu số tiền hiện có của user > tổng tiền đơn hàng thì thực hiện
        if ex_balance >= total:
            # Trừ tiền
            balance = ex_balance - total
            update_user = {"$set": {
                "balance": balance,
            }}
            Login.update_one(found_user, update_user)
            # cập nhật trạng thái đơn hàng chuyển từ false sang True (đã gửi yêu cầu mua hàng)
            update_order = {"$set": {
                "is_ordered": True,
            }}
            Order.update_one(found_order, update_order)
            # Gửi Mail
            gmail = GMail('ngovankhanh108@gmail.com', 'KhanhNgo108')
            html_content = '''Đơn hàng của bạn đã được gửi đi,
                  Tài khoản của bạn sẽ bị trừ đi {{total}}$ tổng chi phí sản phẩm và vận chuyển của đơn hàng này.
                  Cảm ơn bạn đã sử dụng dịch vụ của chúng tôi'''
            html_content1 = html_content.replace("{{total}}", str(total))
            msg = Message(
                'GỬI YÊU CẦU',
                to=found_user['email'],
                html=html_content1
            )
            gmail.send(msg)
            return render_template('user/accepted.html', template=1)
        else:
            return render_template('user/accepted.html', template=0)
    else:
        return 'Order is not found'

# Hiển thị trạng thái đơn hàng đã yêu cầu
@app.route('/order_status')
def orderStatus(): # hàm hiển thị tất cả các đơn hàng đã gửi đi của user
    # tìm tất cả các đơn hàng của user đăng nhập với id bằng phiên id đăng nhập hiện tại, tình trạng order = true(đã gửi yêu cầu),shipper_id = null(đã gửi yêu cầu và chưa shipper nhận đơn)
    all_order = Order.find({'user_id': session['user_id'], 'is_ordered': True, 'shipper_id': "null"})
    # đếm số lượng đơn order tìm thấy trong database
    c = int(all_order.count())
    # nếu > 0(tồn tại 1 đơn trở lên) sẽ thực hiện hiển thị
    if c > 0:
        return render_template('user/order-status.html', all_order=all_order, template=1)
    else: # trả về template thông báo hiện không có đơn nào tồn tại
        return render_template('user/order-status.html', template=0)

# Hiển thị trạng thái đơn hàng đã yêu cầu và có shipper nhận
@app.route('/order_status_receiver')
def orderStatusReceiver(): # hàm hiển thị đơn hàng đã có shipper nhận
    # thực hiện tìm kiếm những đơn có shipper chấp nhận vận chuyển, thông tin vận chuyển lưu vào status
    all_order = Order.find({'user_id': session['user_id'], 'is_ordered': True,"status": "Shipper đã nhận hàng, bắt đầu tiến hành vận chuyển"})
    c = int(all_order.count())
    if c > 0:
        return render_template('user/order-status.html', all_order=all_order, template=1)
    else:
        return render_template('user/order-status.html', template=0)

# Hai nút tìm kiếm tình trạng các đơn hàng của user
@app.route('/user_status_select') 
def user_status_select():
    # hàm phụ để hiện thị 2 nút bấm cho người dùng khi người dùng vào trang các đơn hiện tại
    return render_template('user/user-status-select.html')

# Cập nhật địa chỉ ship đơn hàng
@app.route('/update_address', methods=['GET', 'POST'])
def updateAddress():
    # hàm cập nhật địa chỉ ship hàng trước khi gửi yêu cầu mua hàng, mặc định sẽ là địa chỉ mặc định khi đăng ký tài khoản
    found_order = Order.find_one({'user_id': session['user_id'], 'is_ordered': False, 'status': "Shipper chưa nhận đơn"})
    if request.method == 'GET':
        return render_template('user/update-address.html', found_order=found_order)
    elif request.method == 'POST':
        form = request.form
        address = form['address']
        # thực hiện update địa chỉ ship hàng
        update = {"$set": {
            "address": address,
        }}
        Order.update_one(found_order, update)
    return redirect('/cart')

# User xác nhận đã nhận được hàng
@app.route('/receive_order/<order_id>')
def receive_order(order_id):
    # khi nhận được hàng, user thực hiện bấm nút xác nhận "Đã nhận được hàng."
    # tìm đơn order hiện tại
    found_order = Order.find_one({"_id": ObjectId(order_id)})
    # nếu tìm thấy đơn order
    if found_order is not None:
            # Chuyển trạng thái đơn sang false và cập nhật trạng thái status = "User đã nhận được hàng"
        update1 = {"$set": {
            "status": "User đã nhận hàng",
            "is_ordered": False,
        }}
        Order.update_one(found_order, update1)
        # Add phí vận chuyển cho Shipper
        # tìm thông tin shipper qua id
        id_shipper = found_order['shipper_id']
        found_shipper = Login.find_one({"_id": ObjectId(id_shipper)})
        # thực hiện cộng tiền ship vào tài khoản shipper
        ship_fee = found_order['ship_fee']
        if found_shipper['balance'] == 0:
            update2 = {"$set": {"balance": ship_fee, }}
            Login.update_one(found_shipper, update2)
        else:
            balance = int(found_shipper['balance'])+int(ship_fee)
            update3 = {"$set": {"balance": balance, }}
            Login.update_one(found_shipper, update3)
        # Gửi mail cho Shipper
        gmail = GMail('ngovankhanh108@gmail.com', 'KhanhNgo108')
        html_content = '''Bạn đã vận chuyển thành công đơn hàng đến tay người dùng
                    Bạn sẽ nhận được {{ship_fee}}$ phí ship cho đơn hàng này.
                    Cảm ơn bạn đã tham gia vào hệ thống của chúng tôi'''
        html_content1 = html_content.replace("{{ship_fee}}", str(ship_fee))
        msg = Message(
            'NHẬN ĐƠN',
            to=found_shipper['email'],
            html=html_content1
        )
        gmail.send(msg)
        return redirect('/order_status')
    else:
        return 'Order is not found'

# Hiển thị lịch sử mua hàng
@app.route('/order_history')
def order_history():
    # thực hiện tìm kiếm những đơn hàng mà user đã xác nhận giao hàng thành công
    all_order = Order.find(
        {'user_id': session['user_id'], 'is_ordered': False, 'status': "User đã nhận hàng"})
    if all_order is not None:
        # nếu tồn tại thì hiển thị các đơn hàng
        return render_template('user/order-history.html', all_order=all_order, template=1)
    else:
        # thông báo cho người dùng biết không có đơn hàng nào tồn tại
        return render_template('user/order-history.html', template=0)

# Logout shipper login user
@app.route('/logoutShipper_loginUser')
def logoutShipper_loginUser():
    # hàm thực hiện đăng xuất phiên đăng nhập
    if session['loggedin'] == False:
        return redirect('/login')
    else:
        session['loggedin'] = False
        if session['level_loggedin'] == "1": # thực hiện kiểm tra lần lượt các session level
            del session['admin_id'] # thực hiện xóa session để đăng xuất
            del session['level_loggedin']# thực hiện xóa session để đăng xuất
        elif session['level_loggedin'] == "2":
            del session['user_id']
            del session['level_loggedin']
        elif session['level_loggedin'] == "3":
            del session['shipper_id']
            del session['level_loggedin']
        return redirect('/login')

# Logout admin login user
@app.route('/logoutAdmin_loginUser')
def logoutAdmin_loginUser():
    if session['loggedin'] == False:
        return redirect('/login')
    else:
        session['loggedin'] = False
        if session['level_loggedin'] == "1":
            del session['admin_id']
            del session['level_loggedin']
        elif session['level_loggedin'] == "2":
            del session['user_id']
            del session['level_loggedin']
        elif session['level_loggedin'] == "3":
            del session['shipper_id']
            del session['level_loggedin']
        return redirect('/login')

# Todo SHIPPER
# Chức năng của shipper

# Đăng ký
@app.route('/sign_up_shipper', methods=['GET', 'POST'])
def signUpShipper():
    if request.method == 'GET':
        return render_template('signup/signup-shipper.html')
    elif request.method == 'POST':
        form = request.form
        fullname = form['fullname']
        username = form['username']
        password = form['password']
        # level = form['level']
        email = form['email']
        phone = form['phone']
        address = form['address']
        id_person = form['id_person']
        # balance = form['balance']

        login = Login.find({"username": username})
        c = int(login.count())
        if c > 0:
            return render_template('signup/signup-fail.html', template=0)
        else:
            new_login = {
                "fullname": fullname,
                "username": username,
                "password": password,
                "level": 3,
                "email": email,
                "phone": phone,
                "id_person": id_person,
                "address": address,
                "balance": 0,
            }
            Login.insert_one(new_login)

            found_login = Login.find_one({'username': username})
            gmail = GMail('ngovankhanh108@gmail.com', 'KhanhNgo108')
            html_content = '''
            Chúc mừng bạn đã tạo tài khoản thành công trên hệ thống.
            Cảm ơn quý khách đã sử dụng dịch vụ của chúng tôi!
            '''
            msg = Message('Tạo tài khoản thành công',
                          to=found_login['email'], html=html_content)
            gmail.send(msg)
            return redirect('/login')

# Thông tin cá nhân
@app.route('/shipper_information')
def shipperInformation():
    login = Login.find_one({"_id": ObjectId(session['shipper_id'])})
    if login is not None:
        return render_template('shipper/information.html', found_shipper=login)
    else:
        return 'Login is not found'

# Cập nhật thông tin cá nhân cho Shipper
@app.route('/shipper_update/<id>', methods=['GET', 'POST'])
def updateShipper(id):
    shipper = Login.find_one({"_id": ObjectId(session['shipper_id'])})
    if shipper is not None:
        if request.method == 'GET':
            return render_template('shipper/shipper-update.html', update_shipper=shipper)
        elif request.method == 'POST':
            form = request.form

            password = form['password']

        update = {"$set": {
            "password": password,
        }}
        Login.update_one(shipper, update)
    else:
        return 'Shipper is not found'
    return redirect('/shipper_information')

# Các đơn hàng hiện có cho shipper
@app.route('/ship_request')
def requestShipper():
    # thực hiện tìm kiếm các đơn order user đã gửi và chưa có người nhận
    all_order = Order.find(
        {'is_ordered': True, 'status': "Shipper chưa nhận đơn"})
    c = int(all_order.count())
    if c > 0:
        # trả về trang html các yêu cầu đơn order cho shipper
        return render_template('shipper/ship-request.html', template=1, all_order=all_order)
    else:
        return render_template('shipper/ship-request.html', template=0)

# Hai nút tìm kiếm tình trạng các đơn Shipper đã nhận
@app.route('/ship_status_select')
def ship_status_select():
    return render_template('shipper/ship-status-select.html')

# Chi tiết đơn hàng
@app.route('/detail_request/<order_id>')
def requestDetail(order_id): # hàm nhận id của đơn order để thực hiện
    # thực hiện tìm kiếm đơn order
    found_order = Order.find_one({"_id": ObjectId(order_id)})
    if found_order is not None:
        # tìm kiếm thông tin user từ id_user lưu trong đơn order
        user = Login.find_one({"_id": ObjectId(found_order['user_id'])})
        # thực hiện lấy mảng chứa thông tin các sản phẩm
        order_product = found_order['product_id']
        return render_template('shipper/detail-request.html', found_order=found_order, user=user, order_product=order_product)
    else:
        return 'Order is not found'

# Lịch sử ship đơn hàng
@app.route('/shipped_history')
def historyShipper():
    # những đơn shipper đã giao thành công và user xác nhận đã nhận được hàng
    # thực hiện tìm kiếm tất cả các đơn đã giao của shipper hiện tại
    all_order = Order.find({'shipper_id': session['shipper_id'], 'status': "User đã nhận hàng"})
    # đếm số lượng đơn order
    c = int(all_order.count())
    # nếu tồn tại thì thực hiện truyền dữ liệu sang html và hiển thị
    if c > 0:
        return render_template('shipper/shipped-history.html', all_order=all_order, template=1)
    else:
        return render_template('shipper/shipped-history.html', template=0)

# Shipper nhận đơn hàng
@app.route('/shipper_accepted_order/<order_id>')
def shipper_accepted_order(order_id): # nhận đơn hàng khi user gửi yêu cầu với id đơn order
    # thực hiện tìm kiếm đơn order
    found_order = Order.find_one({"_id": ObjectId(order_id)})
    if found_order is not None:
        # lấy ra id user
        id_user = found_order['user_id']
        # thực hiện tìm kiếm user
        found_user = Login.find_one({"_id": ObjectId(id_user)})
        # Thêm ID của Shipper vào đơn và đổi trạng thái đơn
        update = {"$set": {"shipper_id": session['shipper_id'],
                           "status": "Shipper đã nhận đơn, đang tiến hành lấy hàng"}}
        Order.update_one(found_order, update)
        # Gửi Mail cho User
        gmail = GMail('ngovankhanh108@gmail.com', 'KhanhNgo108')
        html_content = '''Đơn hàng của bạn đã được Shipper chấp nhận giao hàng, vui lòng kiểm tra tình trạng đơn.
                Cảm ơn bạn đã sử dụng dịch vụ của chúng tôi'''
        msg = Message(
            'NHẬN ĐƠN',
            to=found_user['email'],
            html=html_content
        )
        gmail.send(msg)
        return redirect('/ship_request')
    else:
        return 'Order is not found'

# Xem thông tin đơn hiện tại Shipper chấp nhận vận chuyển
@app.route('/ship_status/<status>')
def ship_status(status):
    # những đơn hàng má shipper nhận sẽ có 2 trạng thái
    # trạng thái 1: Shipper đã nhận đơn, đang tiến hành lấy hàng
    # trạng thái 2: Shipper đã nhận hàng, bắt đầu tiến hành vận chuyển
    # thực hiện tìm kiếm status với trạng thái trên
    found_order = Order.find({'shipper_id': session['shipper_id'], 'status': status})
    c = int(found_order.count())
    if c > 0:
        return render_template('shipper/ship-status.html', found_order=found_order, template=1)
    else:
        return render_template('shipper/ship-status.html', template=0)

# Shipper xác nhận đã lấy hàng
@app.route('/product_request/<order_id>')
def product_request(order_id):
    # trạng thái thứ 2 khi shipper đã xác nhận lấy hàng và bắt đầu vận chuyển
    # tìm đơn order hiện tại
    found_order = Order.find_one({"_id": ObjectId(order_id)})
    if found_order is not None:
        # lấy thông tin id user
        id_user = found_order['user_id']
        # tìm thông tin user
        found_user = Login.find_one({"_id": ObjectId(id_user)})
        # thực hiện update trạng thái giao hàng trên database
        update = {
            "$set": {"status": "Shipper đã nhận hàng, bắt đầu tiến hành vận chuyển"}}
        Order.update_one(found_order, update)
        # gửi email thông báo
        gmail = GMail('ngovankhanh108@gmail.com', 'KhanhNgo108')
        msg = Message(
            'LẤY HÀNG', to=found_user['email'], html="Shipper đã lấy được hàng, bắt đầu vận chuyển đến bạn")
        gmail.send(msg)
        return redirect('/shipped_history')
    else:
        return 'Order is not found'


# TODO Admin
# Chức năng của admin
# Information Admin
@app.route('/admin_information')
def adminInformation():
    # thông tin cá nhận 
    login = Login.find_one({"_id": ObjectId(session["admin_id"])})
    if login is not None:
        return render_template('admin/information.html', found_admin=login)
    else:
        return 'Login is not found'

# Update thông tin admin
@app.route('/admin_update/<id>', methods=['GET', 'POST'])
def updateInfoAdmin(id):
    login = Login.find_one({"_id": ObjectId(id)})
    if request.method == 'GET':
        return render_template('admin/admin-update.html', update_admin=login)
    elif request.method == 'POST':
        form = request.form

        password = form['password']

        update_login = {"$set": {
            "password": password,
        }}
        Login.update_one(login, update_login)
        return redirect('/homeLogin')

# Hiển thị sản phẩm
@app.route('/all_product_admin')
def allProduct():
    # tìm kiếm tất cả các sản phẩm có trên database
    all_product = Product.find()
    return render_template('admin/all_product.html', all_product=all_product)

# Xóa sản phẩm
@app.route('/product/delete/<id>')
def deleteProduct(id):
    # thực hiện xóa 1 sản phẩm khi 1 id sản phẩm được truyền vào hàm
    product = Product.find_one({"_id": ObjectId(id)})
    Product.delete_one(product)
    return redirect('/all_product_admin')

# Sửa sản phẩm
@app.route('/all_product/update_product/<id>', methods=['GET', 'POST'])
def updateProduct(id):
    product = Product.find_one({"_id": ObjectId(id)})
    if request.method == 'GET':
        return render_template('admin/update_product.html', product=product)
    elif request.method == 'POST':
        form = request.form

        image = form['image']
        name = form['name']
        description = form['description']
        status = form['status']
        price = form['price']
        brand = form['brand']
        product_type = form['product_type']
        product_gender = form['product_gender']
        product_kids = form['product_kids']

        update_product = {"$set": {
            "image": image,
            "name": name,
            "description": description,
            "price": price,
            "status": status,
            "brand": brand,
            "product_type": product_type,
            "product_gender": product_gender,
            "product_kids": product_kids,
        }}
        Product.update_one(product, update_product)
        return redirect('/all_product_admin')

# Thêm sản phẩm
@app.route('/add_product', methods=['GET', 'POST'])
def addProduct():
    if request.method == 'GET':
        return render_template('admin/add_product.html')
    elif request.method == 'POST':
        form = request.form

        image = form['image']
        name = form['name']
        description = form['description']
        status = form['status']
        price = form['price']
        brand = form['brand']
        product_type = form['product_type']
        product_gender = form['product_gender']
        product_kids = form['product_kids']

        new_product = {
            "image": image,
            "name": name,
            "description": description,
            "status": status,
            "price": price,
            "brand": brand,
            "product_type": product_type,
            "product_gender": product_gender,
            "product_kids": product_kids,
            "view":1,
            "sold_count":1,
        }

        Product.insert_one(new_product)
        return redirect('/all_product_admin')

# Quản lý tài khoản
@app.route('/all_login')
def allLogin():
    # tìm kiếm tất cả các tài khoản hiện có
    all_login = Login.find()
    return render_template('admin/all_login.html', all_login=all_login)

# Xóa tài khoản
@app.route('/all_login/delete/<id>')
def deleteLogin(id):
    # thực hiện xóa tải khoản khi admin thực hiện
    login = Login.find_one({"_id": ObjectId(id)})
    Login.delete_one(login)
    return redirect('/all_login')

# Sửa tài khoản
@app.route('/edit_logins_admin/<id>', methods=['GET', 'POST'])
def editLoginsAdmin(id):
    # sửa các tài khoản hiện có trên database( quyền admin)
    login = Login.find_one({"_id":ObjectId(id)})
    if request.method == "GET":
        return render_template('admin/update_account.html',login = login)
    elif request.method =="POST":
        form = request.form

        update = {"$set":{
            'fullname' : form['fullname'],
            'username' : form['username'],
            'password' : form['password'],
            'level ': form['level'],
            'email' : form['email'],
            'phone' : form['phone'],
            'address' : form['address'],
            'balance' : form['balance'],
        }}

        Login.update_one(login,update)
        return redirect('/all_login')

# Xuất file Excel Account
@app.route('/output_excel')
def output_excel():
    # lấy tất cả tài khoản
    all_login = Login.find()
    # tạo mảng để lưu các trường
    data = []
    i = 0
    for login in all_login:
        list_login = {}
        # lấy giá trị lưu vào biến
        index = i
        fullname = login['fullname']
        username = login['username']
        password = login['password']
        level = login['level']
        email = login['email']
        phone = login['phone']
        balance = login['balance']
        # Key
        list_login['STT'] = index
        list_login['Fullname'] = fullname
        list_login['Username'] = username
        list_login['Password'] = password
        list_login['Level'] = level
        list_login['Email'] = email
        list_login['Phone'] = phone
        list_login['Balance'] = balance
        i = i + 1
        #
        # thêm 1 list mới vào mảng
        data.append(list_login)
    # thực hiện xuất file excel
    pyexcel.save_as(records=data, dest_file_name="Logins.xlsx")
    return "Done"

# Xuất file Excel products
@app.route('/output_product_excel')
def output_product_excel():
    all_product = Product.find()
    data = []
    i = 0
    for product in all_product:
        list_product = {}
        #
        index = i
        _id = product['_id']
        image = product['image']
        name = product['name']
        description = product['description']
        price = product['price']
        status = product['status']
        brand = product['brand']
        product_type = product['product_type']
        product_gender = product['product_gender']
        #
        list_product['STT'] = index
        list_product['ID'] = _id
        list_product['Name'] = name
        list_product['Description'] = description
        list_product['Price'] = price
        list_product['Status'] = status
        list_product['Brand'] = brand
        list_product['Type'] = product_type
        list_product['Gender'] = product_gender
        i = i + 1
        #
        data.append(list_product)
    pyexcel.save_as(records=data, dest_file_name="Products.xlsx")
    return "Done"

#  Xuất file Order
@app.route('/output_orders_excel')
def output_order_excel():
    all_order = Order.find()
    data = []
    i = 0
    for order in all_order:
        list_order = {}
        #
        index = i
        _id = order['_id']
        user_id = order['user_id']
        address = order['address']
        order_time = order['order_time']
        order_fee = order['order_fee']
        ship_fee = order['ship_fee']
        is_ordered = order['is_ordered']
        status = order['status']
        shipper_id = order['shipper_id']
        #
        list_order['STT'] = index
        list_order['ID'] = _id
        list_order['User ID'] = user_id
        list_order['Address'] = address
        list_order['Order Time'] = order_time
        list_order['Order Fee'] = order_fee
        list_order['Ship Fee'] = ship_fee
        list_order['Is Ordered'] = is_ordered
        list_order['Status'] = status
        list_order['Shipper ID'] = shipper_id
        i = i + 1

        #
        data.append(list_order)
    pyexcel.save_as(records=data, dest_file_name="Orders.xlsx")

    return "Done"

# Xuất hóa đơn ra Word
@app.route('/output_bill_word/<id>')
def output_bill_word(id):
    # tìm đơn order muốn xuất
    order = Order.find_one({"_id":ObjectId(id)})
    # lấy thông tin các sản phẩm
    all_product =order['product_id']
    # tìm thông tin user
    user = Login.find_one({"_id":ObjectId(order['user_id'])})
    # khởi tạo đối tượng
    document = Document()
    # add tiêu đều
    document.add_heading('Hóa đơn bán hàng', 0)
    # lấy ngày giờ order
    date = order['order_time']
    day = str(date.day)
    month = str(date.month)
    year = str(date.year)
    # thêm ngày giờ vào hóa đơn
    document.add_heading('Ngày ' + day+"/"+month+"/"+year, 0)
    # lấy tên khách hàng từ database
    document.add_heading("Tên khách hàng:  " + user['fullname'], level=1)
    # duyết mảng sản phẩm để lấy thông tin sản phẩm
    index = 1
    for i in all_product:
        name = i['name']
        price = str(i['price'])
        s = "Sản phẩm " + str(index) + " :  " + name + "             " + price+"$"
        document.add_paragraph(s, style='Intense Quote')
        index += 1
    # lấy thông tin địa chỉ từ đơn order
    document.add_paragraph("Địa chỉ:  " + order['address'], style='Intense Quote')
    # lấy phí ship từ đơn hàng
    document.add_paragraph(
        "Phí ship:  " + str(order['ship_fee']), style='Intense Quote')
    # tổng tiền hàng
    document.add_paragraph("Tổng tiền hàng : " +
                        str(order['order_fee']), style='Intense Quote')
    # tổng tiền của dơn order phí ship + tiền hàng
    totalPrice = int(order['ship_fee']) + int(order['order_fee'])
    document.add_paragraph("Tổng thanh toán:  " +
                        str(totalPrice), style='Intense Quote')

    document.add_page_break()
    # xuất file word
    document.save('HoaDonWord.docx')
    return "done"

# TODO
@app.route('/all_order')
def allOrder():
    all_order = Order.find()
    return render_template('admin/all_order.html',all_order = all_order)

if __name__ == '__main__':
    app.run(debug=True)
